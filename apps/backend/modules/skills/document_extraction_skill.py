import os
import logging
from livekit.agents import llm
from modules.skills.base_skill import BaseSkill
from modules.filesystem.document_extractor import DocumentExtractor

logger = logging.getLogger("JARVIS.Skills.DocumentExtractionSkill")

class DocumentExtractionSkill(BaseSkill):
    """
    Skill to extract text and structured content from documents like PDF, DOCX, XLSX, and images,
    and convert or manipulate document content.
    """
    def __init__(self, memory=None, security=None, room=None, verification=None, **kwargs):
        super().__init__(memory=memory, security=security, room=room, verification=verification)
        self.extractor = DocumentExtractor()

    @llm.function_tool(
        description="Extract text, tables, and metadata from a document (PDF, DOCX, XLSX, PPTX, CSV, JSON, YAML, HTML, XML, or image) at the specified path"
    )
    async def extract_document(self, file_path: str, ocr_embedded_images: bool = False) -> str:
        """Extracts content from a file and returns it in detail."""
        from modules.security.manager import SecurityManager
        sec = self.security if self.security else SecurityManager()
        
        act_path = os.path.normpath(os.path.abspath(file_path))
        if not sec.is_safe_path(act_path):
            return "Error: Security Policy blocks reading protected system path."
            
        async def _do_extract():
            result = self.extractor.extract(act_path, ocr_embedded_images=ocr_embedded_images)
            if not result.success:
                return f"Error extracting document: {result.error or 'Unknown error'}"
            
            output = f"Extraction Successful (Method: {result.method})\n"
            output += f"Metadata: {result.metadata}\n\n"
            if result.text:
                output += f"--- Extracted Text ---\n{result.text}\n"
            else:
                output += "No text content found.\n"
                
            return output

        return await self.safe_execute(
            _do_extract,
            confirmation_category="read",
            confirmation_action=f"extract document from {file_path}",
            confirmed=True,
            error_msg="Failed to extract document"
        )

    @llm.function_tool(
        description="Convert a document from a source format to a target format. Supports CSV <-> XLSX, XLSX <-> JSON, JSON <-> CSV, CSV <-> JSON, DOCX <-> TXT, TXT <-> DOCX."
    )
    async def convert_document(self, src_path: str, dest_path: str) -> str:
        """Converts between structured and unstructured formats."""
        from modules.security.manager import SecurityManager
        sec = self.security if self.security else SecurityManager()
        
        src_abs = os.path.normpath(os.path.abspath(src_path))
        dest_abs = os.path.normpath(os.path.abspath(dest_path))
        
        if not sec.is_safe_path(src_abs) or not sec.is_safe_path(dest_abs):
            return "Error: Security Policy blocks conversion between these paths."
            
        async def _do_convert():
            src_ext = os.path.splitext(src_abs)[1].lower()
            dest_ext = os.path.splitext(dest_abs)[1].lower()
            
            import pandas as pd
            import docx
            import json
            
            os.makedirs(os.path.dirname(dest_abs), exist_ok=True)
            
            if src_ext == ".csv" and dest_ext == ".xlsx":
                df = pd.read_csv(src_abs)
                df.to_excel(dest_abs, index=False)
                return f"Successfully converted CSV to Excel: {dest_path}"
                
            elif src_ext == ".xlsx" and dest_ext == ".csv":
                with pd.ExcelFile(src_abs) as xlsx:
                    sheet = xlsx.sheet_names[0]
                    df = pd.read_excel(xlsx, sheet_name=sheet)
                df.to_csv(dest_abs, index=False)
                return f"Successfully converted Excel sheet '{sheet}' to CSV: {dest_path}"
                
            elif src_ext == ".xlsx" and dest_ext == ".json":
                with pd.ExcelFile(src_abs) as xlsx:
                    sheet = xlsx.sheet_names[0]
                    df = pd.read_excel(xlsx, sheet_name=sheet)
                df = df.fillna("")
                df.to_json(dest_abs, orient="records", indent=2)
                return f"Successfully converted Excel sheet '{sheet}' to JSON: {dest_path}"
                
            elif src_ext == ".json" and dest_ext == ".xlsx":
                with open(src_abs, "r", encoding="utf-8") as f:
                    data = json.load(f)
                if isinstance(data, list):
                    df = pd.DataFrame(data)
                elif isinstance(data, dict):
                    list_key = next((k for k, v in data.items() if isinstance(v, list)), None)
                    if list_key:
                        df = pd.DataFrame(data[list_key])
                    else:
                        df = pd.DataFrame([data])
                else:
                    return "Error: JSON structure is not suitable for tabular conversion (must be array or contain arrays)."
                df.to_excel(dest_abs, index=False)
                return f"Successfully converted JSON to Excel: {dest_path}"
                
            elif src_ext == ".json" and dest_ext == ".csv":
                with open(src_abs, "r", encoding="utf-8") as f:
                    data = json.load(f)
                if isinstance(data, list):
                    df = pd.DataFrame(data)
                elif isinstance(data, dict):
                    list_key = next((k for k, v in data.items() if isinstance(v, list)), None)
                    if list_key:
                        df = pd.DataFrame(data[list_key])
                    else:
                        df = pd.DataFrame([data])
                else:
                    return "Error: JSON structure is not suitable for tabular conversion (must be array or contain arrays)."
                df.to_csv(dest_abs, index=False)
                return f"Successfully converted JSON to CSV: {dest_path}"
                
            elif src_ext == ".csv" and dest_ext == ".json":
                df = pd.read_csv(src_abs)
                df = df.fillna("")
                df.to_json(dest_abs, orient="records", indent=2)
                return f"Successfully converted CSV to JSON: {dest_path}"
                
            elif src_ext == ".docx" and dest_ext == ".txt":
                doc = docx.Document(src_abs)
                paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
                with open(dest_abs, "w", encoding="utf-8") as f:
                    f.write("\n".join(paragraphs))
                return f"Successfully converted DOCX to TXT: {dest_path}"
                
            elif src_ext == ".txt" and dest_ext == ".docx":
                with open(src_abs, "r", encoding="utf-8") as f:
                    content = f.read()
                doc = docx.Document()
                for line in content.split("\n"):
                    if line.strip():
                        doc.add_paragraph(line)
                doc.save(dest_abs)
                return f"Successfully converted TXT to DOCX: {dest_path}"
                
            else:
                return f"Error: Conversion from {src_ext} to {dest_ext} is not supported."

        return await self.safe_execute(
            _do_convert,
            confirmation_category="open",
            confirmation_action=f"convert {src_path} to {dest_path}",
            confirmed=True,
            error_msg="Failed to convert document"
        )

    @llm.function_tool(
        description="Modify structured contents of a document (append row/cells to CSV/XLSX, add text/tables to DOCX, update JSON/YAML keys)"
    )
    async def modify_document_data(self, file_path: str, modification_type: str, updates: dict) -> str:
        """Modifies file contents based on format."""
        from modules.security.manager import SecurityManager
        sec = self.security if self.security else SecurityManager()
        
        act_path = os.path.normpath(os.path.abspath(file_path))
        if not sec.is_safe_path(act_path):
            return "Error: Security Policy blocks modification of this path."
            
        async def _do_modify():
            ext = os.path.splitext(act_path)[1].lower()
            
            import pandas as pd
            import docx
            import json
            
            if ext == ".csv":
                df = pd.read_csv(act_path)
                if modification_type == "append_row":
                    row_data = updates.get("row")
                    if isinstance(row_data, dict):
                        df = pd.concat([df, pd.DataFrame([row_data])], ignore_index=True)
                    elif isinstance(row_data, list):
                        if len(row_data) == len(df.columns):
                            df.loc[len(df)] = row_data
                        else:
                            return f"Error: Row length ({len(row_data)}) does not match CSV columns count ({len(df.columns)})."
                    else:
                        return "Error: 'row' parameter must be a list of values or a dictionary of column-values."
                elif modification_type == "update_cell":
                    row = int(updates.get("row"))
                    col = updates.get("col")
                    val = updates.get("value")
                    if isinstance(col, str) and col in df.columns:
                        df.at[row, col] = val
                    elif isinstance(col, (int, float)):
                        df.iloc[row, int(col)] = val
                    else:
                        return f"Error: Column '{col}' is not valid."
                else:
                    return f"Error: Modification type '{modification_type}' is not supported for CSV."
                    
                df.to_csv(act_path, index=False)
                return f"Successfully updated CSV document: {file_path}"
                
            elif ext == ".xlsx":
                with pd.ExcelFile(act_path) as xlsx:
                    sheet_names = xlsx.sheet_names
                    all_sheets = {name: pd.read_excel(xlsx, sheet_name=name) for name in sheet_names}
                
                sheet = updates.get("sheet", sheet_names[0])
                if sheet not in all_sheets:
                    all_sheets[sheet] = pd.DataFrame()
                    
                df = all_sheets[sheet]
                if modification_type == "append_row":
                    row_data = updates.get("row")
                    if isinstance(row_data, dict):
                        df = pd.concat([df, pd.DataFrame([row_data])], ignore_index=True)
                    elif isinstance(row_data, list):
                        if df.empty:
                            df = pd.DataFrame([row_data])
                        elif len(row_data) == len(df.columns):
                            df.loc[len(df)] = row_data
                        else:
                            return f"Error: Row length does not match Excel sheet columns count."
                    else:
                        return "Error: 'row' must be list or dict."
                elif modification_type == "update_cell":
                    row = int(updates.get("row"))
                    col = updates.get("col")
                    val = updates.get("value")
                    if df.empty:
                        return "Error: Cannot update cell in an empty sheet."
                    if isinstance(col, str) and col in df.columns:
                        df.at[row, col] = val
                    elif isinstance(col, (int, float)):
                        df.iloc[row, int(col)] = val
                    else:
                        return f"Error: Column '{col}' is not valid."
                else:
                    return f"Error: Modification type '{modification_type}' is not supported for XLSX."
                
                all_sheets[sheet] = df
                with pd.ExcelWriter(act_path, engine="openpyxl") as writer:
                    for s_name, s_df in all_sheets.items():
                        s_df.to_excel(writer, sheet_name=s_name, index=False)
                return f"Successfully updated Excel sheet '{sheet}' at: {file_path}"
                
            elif ext == ".docx":
                doc = docx.Document(act_path)
                if modification_type == "append_paragraph":
                    text = updates.get("text", "")
                    doc.add_paragraph(text)
                elif modification_type == "append_table":
                    matrix = updates.get("table", [])
                    if not matrix:
                        return "Error: 'table' matrix cannot be empty."
                    rows_cnt = len(matrix)
                    cols_cnt = len(matrix[0])
                    table = doc.add_table(rows=rows_cnt, cols=cols_cnt)
                    for r_idx, row in enumerate(matrix):
                        for c_idx, val in enumerate(row):
                            table.cell(r_idx, c_idx).text = str(val)
                else:
                    return f"Error: Modification type '{modification_type}' is not supported for DOCX."
                doc.save(act_path)
                return f"Successfully updated DOCX document: {file_path}"
                
            elif ext == ".json":
                with open(act_path, "r", encoding="utf-8") as f:
                    data = json.load(f)
                    
                if modification_type == "set_key":
                    key = updates.get("key")
                    val = updates.get("value")
                    if isinstance(data, dict):
                        data[key] = val
                    else:
                        return "Error: JSON root is not an object/dictionary."
                elif modification_type == "append_list":
                    val = updates.get("value")
                    if isinstance(data, list):
                        data.append(val)
                    else:
                        return "Error: JSON root is not a list/array."
                else:
                    return f"Error: Modification type '{modification_type}' not supported for JSON."
                    
                with open(act_path, "w", encoding="utf-8") as f:
                    json.dump(data, f, indent=2)
                return f"Successfully updated JSON document: {file_path}"
                
            else:
                return f"Error: Modifying {ext} documents is not natively supported."

        return await self.safe_execute(
            _do_modify,
            confirmation_category="open",
            confirmation_action=f"modify document {file_path}",
            confirmed=True,
            error_msg="Failed to modify document"
        )

    @llm.function_tool(
        description="Filter or query tabular data in a CSV or Excel document using a pandas query string (e.g. `Age > 30` or `Department == 'Sales'`)"
    )
    async def query_document_data(self, file_path: str, pandas_query: str, sheet_name: str = None) -> str:
        """Filters table data based on pandas query execution."""
        from modules.security.manager import SecurityManager
        sec = self.security if self.security else SecurityManager()
        
        act_path = os.path.normpath(os.path.abspath(file_path))
        if not sec.is_safe_path(act_path):
            return "Error: Security Policy blocks access to this path."
            
        async def _do_query():
            ext = os.path.splitext(act_path)[1].lower()
            import pandas as pd
            
            if ext == ".csv":
                df = pd.read_csv(act_path)
            elif ext == ".xlsx":
                with pd.ExcelFile(act_path) as xlsx:
                    sheet = sheet_name or xlsx.sheet_names[0]
                    df = pd.read_excel(xlsx, sheet_name=sheet)
            else:
                return "Error: Only CSV and XLSX files are supported for semantic tabular queries."
                
            df = df.fillna("")
            
            try:
                filtered_df = df.query(pandas_query)
            except Exception as query_ex:
                return f"Error evaluating query '{pandas_query}': {query_ex}"
                
            if filtered_df.empty:
                return "Query returned 0 matching rows."
                
            headers = " | ".join(str(h) for h in filtered_df.columns)
            markdown_text = f"Query Results (found {len(filtered_df)} rows):\n\n"
            markdown_text += headers + "\n"
            markdown_text += "-|-".join("-" * len(str(h)) for h in filtered_df.columns) + "\n"
            for _, row in filtered_df.iterrows():
                markdown_text += " | ".join(str(val) for val in row.values) + "\n"
                
            return markdown_text.strip()

        return await self.safe_execute(
            _do_query,
            confirmation_category="read",
            confirmation_action=f"query document {file_path} with: {pandas_query}",
            confirmed=True,
            error_msg="Failed to query document"
        )

    @llm.function_tool(
        description="Get the column names, data types, sheet names, row count, and first 3 sample rows from a CSV or Excel spreadsheet"
    )
    async def get_tabular_schema(self, file_path: str, sheet_name: str = None) -> str:
        """Reads schema metadata and shapes from structured table documents."""
        from modules.security.manager import SecurityManager
        sec = self.security if self.security else SecurityManager()
        
        act_path = os.path.normpath(os.path.abspath(file_path))
        if not sec.is_safe_path(act_path):
            return "Error: Security Policy blocks access to this path."
            
        async def _do_schema():
            ext = os.path.splitext(act_path)[1].lower()
            import pandas as pd
            
            sheet_names = []
            if ext == ".csv":
                df = pd.read_csv(act_path)
            elif ext == ".xlsx":
                with pd.ExcelFile(act_path) as xlsx:
                    sheet_names = xlsx.sheet_names
                    sheet = sheet_name or sheet_names[0]
                    df = pd.read_excel(xlsx, sheet_name=sheet)
            else:
                return "Error: Only CSV and XLSX files are supported for schema retrieval."
                
            df = df.fillna("")
            total_rows = len(df)
            columns = df.columns.tolist()
            dtypes = {str(col): str(df[col].dtype) for col in df.columns}
            
            output = f"File: {os.path.basename(file_path)}\n"
            if ext == ".xlsx":
                output += f"Sheet: {sheet} (All sheets in file: {sheet_names})\n"
            output += f"Total Row Count: {total_rows}\n"
            output += f"Columns & Types:\n"
            for col in columns:
                output += f"  - {col}: {dtypes[col]}\n"
                
            output += "\nSample Rows (First 3 rows):\n"
            sample_df = df.head(3)
            headers = " | ".join(str(h) for h in sample_df.columns)
            output += headers + "\n"
            output += "-|-".join("-" * len(str(h)) for h in sample_df.columns) + "\n"
            for _, row in sample_df.iterrows():
                output += " | ".join(str(val) for val in row.values) + "\n"
                
            return output.strip()

        return await self.safe_execute(
            _do_schema,
            confirmation_category="read",
            confirmation_action=f"read schema of document {file_path}",
            confirmed=True,
            error_msg="Failed to retrieve document schema"
        )

    @llm.function_tool(
        description="Read a specific page/slice of rows from a CSV or Excel document (page indexing starts at 1)"
    )
    async def read_tabular_page(self, file_path: str, page: int = 1, page_size: int = 25, sheet_name: str = None) -> str:
        """Reads a paginated slice of rows from structured table documents to prevent truncation."""
        from modules.security.manager import SecurityManager
        sec = self.security if self.security else SecurityManager()
        
        act_path = os.path.normpath(os.path.abspath(file_path))
        if not sec.is_safe_path(act_path):
            return "Error: Security Policy blocks access to this path."
            
        async def _do_page():
            ext = os.path.splitext(act_path)[1].lower()
            import pandas as pd
            
            if ext == ".csv":
                df = pd.read_csv(act_path)
            elif ext == ".xlsx":
                with pd.ExcelFile(act_path) as xlsx:
                    sheet = sheet_name or xlsx.sheet_names[0]
                    df = pd.read_excel(xlsx, sheet_name=sheet)
            else:
                return "Error: Only CSV and XLSX files are supported for paging."
                
            df = df.fillna("")
            total_rows = len(df)
            
            start_row = (page - 1) * page_size
            end_row = start_row + page_size
            
            if start_row >= total_rows or start_row < 0:
                return f"Error: Page {page} is out of bounds (total rows: {total_rows}, page size: {page_size})."
                
            sliced_df = df.iloc[start_row:end_row]
            
            output = f"File: {os.path.basename(file_path)}"
            if ext == ".xlsx":
                output += f" (Sheet: {sheet})"
            output += f"\nRows {start_row + 1} to {min(end_row, total_rows)} of {total_rows} (Page {page} of {(total_rows + page_size - 1) // page_size})\n\n"
            
            headers = " | ".join(str(h) for h in sliced_df.columns)
            output += headers + "\n"
            output += "-|-".join("-" * len(str(h)) for h in sliced_df.columns) + "\n"
            for idx, row in sliced_df.iterrows():
                row_vals = [str(val) for val in row.values]
                output += " | ".join(row_vals) + "\n"
                
            return output.strip()

        return await self.safe_execute(
            _do_page,
            confirmation_category="read",
            confirmation_action=f"read page {page} of document {file_path}",
            confirmed=True,
            error_msg="Failed to read paginated document content"
        )
